import re
import itertools
from io import BytesIO
from typing import List

from docx import Document
from docx.text.run import Run
from tqdm import tqdm

from .gcp_translate import translate_batch

MAX_CHUNK = 4500
NUM_RE   = re.compile(r"\d[\d\s.,]*")

def _runs(doc: Document) -> List[Run]:
    runs = []
    runs.extend([r for p in doc.paragraphs for r in p.runs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    runs.extend(p.runs)
    return runs

def _bucketize(runs: List[Run]) -> List[List[Run]]:
    bucket, buckets, current = [], [], 0
    for r in runs:
        if current + len(r.text) > MAX_CHUNK and bucket:
            buckets.append(bucket)
            bucket, current = [], 0
        bucket.append(r)
        current += len(r.text)
    if bucket:
        buckets.append(bucket)
    return buckets

def translate_docx(file, target_lang: str, gcp_project: str) -> BytesIO:
    doc = Document(file)
    all_runs = _runs(doc)

    placeholders, sliced_texts = [], []
    for r in all_runs:
        text, numbers = NUM_RE.sub("{NUM}", r.text), NUM_RE.findall(r.text)
        placeholders.append(numbers)
        sliced_texts.append(text)

    translated = []
    for bucket in tqdm(_bucketize(all_runs), desc="Translating"):
        texts_to_translate = []
        map_back = []

        for run in bucket:
            clean = NUM_RE.sub("{NUM}", run.text)
            if clean.strip():
                map_back.append(len(translated))
                texts_to_translate.append(clean)
                translated.append(None)
            else:
                translated.append(run.text)

        if texts_to_translate:
            result = translate_batch(
                texts_to_translate,
                target_lang,
                project_id=gcp_project,
            )
            for idx, tr in zip(map_back, result):
                translated[idx] = tr

    for run, tr, nums in zip(all_runs, translated, placeholders):
        for n in nums:
            tr = tr.replace("{NUM}", n, 1)
        run.text = tr

    out_buf = BytesIO()
    doc.save(out_buf)
    out_buf.seek(0)
    return out_buf
